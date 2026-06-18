"""
gen_doc.py – Tạo file .docx tài liệu đầy đủ cho Discord Moderation Bot
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Màu sắc ─────────────────────────────────────────────────────────────────
C_DARK_BG   = RGBColor(0x2C, 0x2F, 0x33)   # Discord dark
C_BLUE      = RGBColor(0x58, 0x65, 0xF2)   # Discord blurple
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT     = RGBColor(0xF5, 0xF5, 0xF5)
C_HEADING   = RGBColor(0x23, 0x27, 0x2A)
C_GREEN     = RGBColor(0x2E, 0xCC, 0x71)
C_RED       = RGBColor(0xE7, 0x4C, 0x3C)
C_ORANGE    = RGBColor(0xF3, 0x9C, 0x12)
C_GRAY      = RGBColor(0x99, 0xAA, 0xB5)
C_CODE_BG   = RGBColor(0xEE, 0xEE, 0xEE)

doc = Document()

# ─── Trang & margin ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Style helpers ───────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading(text: str, level: int = 1, color: RGBColor = C_HEADING):
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(20)
        run.font.color.rgb = C_BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = C_HEADING
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        # đường kẻ dưới tiêu đề
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "5865F2")
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x50, 0x55, 0x60)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    return p


def add_para(text: str, size: int = 10, color: RGBColor = C_HEADING,
             bold: bool = False, italic: bool = False, space_after: int = 4):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.bold   = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code(text: str):
    p   = doc.add_paragraph()
    run = p.add_run(f"  {text}")
    run.font.name      = "Courier New"
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(0x23, 0x27, 0x2A)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    # nền xám nhạt
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EEEEEE")
    pPr.append(shd)
    return p


def add_bullet(text: str, indent: int = 0):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = C_HEADING
    p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_table(headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None,
              header_bg: str = "5865F2"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, "5865F2")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9.5)
        run.font.color.rgb = C_WHITE
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # Data rows
    for ri, row in enumerate(rows):
        drow = table.rows[ri + 1]
        bg   = "F9F9F9" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "E0E0E0")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            # inline code style cho cột lệnh (cột 0)
            if ci == 0 and val.startswith("/"):
                run = p.add_run(val)
                run.font.name      = "Courier New"
                run.font.size      = Pt(9)
                run.font.color.rgb = RGBColor(0x58, 0x65, 0xF2)
                run.bold = True
            else:
                run = p.add_run(val)
                run.font.size      = Pt(9.5)
                run.font.color.rgb = C_HEADING
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    # Độ rộng cột
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_note(text: str, bg: str = "FFF8E1", border: str = "F39C12"):
    """Hộp ghi chú màu."""
    table = doc.add_table(rows=1, cols=1)
    cell  = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_borders(cell, border)
    p   = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size      = Pt(9.5)
    run.font.color.rgb = RGBColor(0x50, 0x40, 0x00) if bg == "FFF8E1" else C_HEADING
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════════════════════════════════
#  TRANG BÌA
# ════════════════════════════════════════════════════════════════════════════

# Banner màu Discord
banner_tbl = doc.add_table(rows=1, cols=1)
banner_cell = banner_tbl.rows[0].cells[0]
set_cell_bg(banner_cell, "5865F2")
bp = banner_cell.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = bp.add_run("\n🛡️  DISCORD MODERATION BOT\n")
r1.font.size      = Pt(24)
r1.font.color.rgb = C_WHITE
r1.bold = True
r2 = bp.add_run("Tài liệu Hướng Dẫn – Toàn Bộ Lệnh & Cách Hoạt Động\n")
r2.font.size      = Pt(12)
r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xFF)
bp.paragraph_format.space_before = Pt(10)
bp.paragraph_format.space_after  = Pt(10)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
mr = meta.add_run(f"Phiên bản: 1.0  |  Ngày tạo: {datetime.date.today().strftime('%d/%m/%Y')}")
mr.font.size      = Pt(9)
mr.font.color.rgb = C_GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  1. TỔNG QUAN
# ════════════════════════════════════════════════════════════════════════════

add_heading("1. Tổng Quan", 1)
add_para(
    "Bot Discord Moderation được xây dựng bằng Python (discord.py 2.x), "
    "tự động kiểm duyệt nội dung và hỗ trợ Staff thực thi nội quy server. "
    "Mọi cấu hình được lưu trong data/config.json và có thể thay đổi trực tiếp "
    "trong Discord qua lệnh /config mà không cần khởi động lại bot.",
    size=10
)

add_heading("Công nghệ", 2)
add_table(
    ["Thành phần", "Chi tiết"],
    [
        ["Ngôn ngữ",    "Python 3.10+"],
        ["Framework",   "discord.py 2.x (Slash + Prefix Commands)"],
        ["Lưu trữ",     "JSON local (data/warnings.json, data/cases.json, data/config.json)"],
        ["Cấu hình",    "Runtime – thay đổi qua /config, không cần .env ngoài TOKEN"],
    ],
    col_widths=[4.5, 11]
)

add_heading("Cấu Trúc File", 2)
add_code("discord-bot/")
add_code("├── bot.py                  ← Entry point, setup wizard")
add_code("├── config.py               ← ConfigManager singleton + defaults")
add_code("├── utils.py                ← DB helpers, embed builder, permission check")
add_code("├── .env                    ← Chỉ chứa DISCORD_TOKEN")
add_code("├── cogs/")
add_code("│   ├── automod.py          ← Kiểm duyệt tự động (8 loại)")
add_code("│   ├── moderation.py       ← Lệnh mod thủ công (9 lệnh)")
add_code("│   ├── rules.py            ← Hiển thị nội quy (10 mục)")
add_code("│   ├── appeals.py          ← Hệ thống kháng cáo (Modal UI)")
add_code("│   ├── logging_events.py   ← Ghi log sự kiện server")
add_code("│   └── config_cmd.py       ← Lệnh /config cho Admin")
add_code("└── data/")
add_code("    ├── config.json         ← Cấu hình runtime (tự tạo)")
add_code("    ├── warnings.json       ← Dữ liệu cảnh cáo thành viên")
add_code("    └── cases.json          ← Lịch sử các case xử lý")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  2. HƯỚNG DẪN CÀI ĐẶT
# ════════════════════════════════════════════════════════════════════════════

add_heading("2. Hướng Dẫn Cài Đặt", 1)

add_heading("Bước 1 – Tạo Bot Discord", 2)
add_bullet("Vào https://discord.com/developers/applications")
add_bullet("Nhấn New Application → đặt tên → vào tab Bot → Add Bot")
add_bullet("Sao chép Token (giữ bí mật)")
add_bullet("Bật Server Members Intent và Message Content Intent")

add_heading("Bước 2 – Cấu hình .env", 2)
add_para("Tạo file .env trong thư mục gốc, chỉ cần 1 dòng:", size=10)
add_code("DISCORD_TOKEN=your_bot_token_here")
add_note("⚠️  Đây là thông tin duy nhất bắt buộc phải điền. Mọi cấu hình khác "
         "(channel, role, ngưỡng spam...) đều có thể đặt qua lệnh /config trong Discord.")

add_heading("Bước 3 – Mời Bot vào Server", 2)
add_bullet("Vào Developer Portal → OAuth2 → URL Generator")
add_bullet("Scopes: bot + applications.commands")
add_bullet("Bot Permissions: Administrator (hoặc chọn từng quyền)")
add_bullet("Mở link được tạo, chọn server, xác nhận")

add_heading("Bước 4 – Chạy Bot", 2)
add_code("pip install -r requirements.txt")
add_code("python bot.py")
add_para("Bot sẽ tự động DM cho Owner server hướng dẫn cấu hình kênh và role.", size=10, italic=True, color=C_GRAY)

add_heading("Bước 5 – Cấu hình trong Discord (lần đầu)", 2)
add_table(
    ["Lệnh", "Ý nghĩa"],
    [
        ["/config channel mod_log #kênh",   "Kênh ghi log hành động mod"],
        ["/config channel rules #kênh",     "Kênh đăng nội quy chính thức"],
        ["/config channel appeal #kênh",    "Kênh nhận kháng cáo"],
        ["/config role moderator @role",    "Role Moderator"],
        ["/config role admin @role",        "Role Admin"],
        ["/config role muted @role",        "Role Muted (tạo thủ công nếu chưa có)"],
    ],
    col_widths=[7, 8.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  3. HỆ THỐNG QUYỀN HẠN
# ════════════════════════════════════════════════════════════════════════════

add_heading("3. Hệ Thống Quyền Hạn", 1)
add_para("Bot kiểm tra quyền theo 2 tầng, hoạt động ngay cả khi chưa cấu hình role:", size=10)

add_heading("Tầng 1 – Fallback (hoạt động ngay)", 2)
add_para("Ai có quyền Administrator trong Discord Settings → Roles được coi là Admin/Mod:", size=10)
add_bullet("✅ Owner server")
add_bullet("✅ Thành viên có tick Administrator trong Server Settings")

add_heading("Tầng 2 – Sau khi cấu hình /config role", 2)
add_table(
    ["Role Config", "Quyền truy cập"],
    [
        ["admin_role_id",  "Toàn bộ lệnh Admin + Staff + /config"],
        ["mod_role_id",    "Lệnh Staff: warn, mute, kick, ban, purge, warns..."],
        ["muted_role_id",  "Không có quyền đặc biệt – dùng cho Timeout thủ công"],
    ],
    col_widths=[5, 10.5]
)

add_heading("Sơ Đồ Kiểm Tra Quyền", 2)
add_code("Người dùng gõ lệnh Staff")
add_code("        ↓")
add_code("Có role Admin ID?  ──Yes──→  ✅ Cho phép")
add_code("        ↓ No")
add_code("Có role Mod ID?    ──Yes──→  ✅ Cho phép")
add_code("        ↓ No")
add_code("Có perm Admin?     ──Yes──→  ✅ Cho phép")
add_code("        ↓ No")
add_code("                   ❌ Từ chối (thông báo ẩn chỉ mình thấy)")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  4. DANH SÁCH LỆNH
# ════════════════════════════════════════════════════════════════════════════

add_heading("4. Danh Sách Lệnh", 1)

# ── 4.1 Thành viên ──────────────────────────────────────────────────────────
add_heading("4.1 Lệnh Dành Cho Tất Cả Thành Viên", 2)
add_table(
    ["Lệnh", "Slash", "Prefix", "Mô tả"],
    [
        ["/help",        "❌",  "✅ !help",   "Xem danh sách lệnh và hướng dẫn"],
        ["/rules",       "✅",  "✅",         "Hiển thị toàn bộ nội quy ngay tại kênh hiện tại"],
        ["/appeal",      "✅",  "✅",         "Mở form kháng cáo (Discord Modal UI)"],
        ["/case [id]",   "✅",  "✅",         "Xem thông tin một case theo ID số"],
    ],
    col_widths=[4, 1.5, 2, 8]
)

# ── 4.2 Staff ────────────────────────────────────────────────────────────────
add_heading("4.2 Lệnh Dành Cho Staff (Mod / Admin)", 2)
add_table(
    ["Lệnh", "Tham số", "Mô tả"],
    [
        ["/warn",       "@user [lý do]",            "Cảnh cáo thành viên. Tự động leo thang hình phạt theo số warn"],
        ["/mute",       "@user [phút] [lý do]",     "Timeout (mute) thành viên. Mặc định 60 phút"],
        ["/unmute",     "@user",                    "Bỏ timeout cho thành viên"],
        ["/kick",       "@user [lý do]",            "Kick thành viên ra khỏi server"],
        ["/ban",        "@user [lý do]",            "Ban vĩnh viễn thành viên"],
        ["/unban",      "[user_id] [lý do]",        "Bỏ ban theo ID người dùng"],
        ["/warns",      "@user",                    "Xem lịch sử cảnh cáo của thành viên"],
        ["/clearwarns", "@user",                    "Xóa toàn bộ cảnh cáo của thành viên"],
        ["/purge",      "[số lượng]",               "Xóa hàng loạt tin nhắn trong kênh (tối đa 100)"],
        ["/propose",    "@user [action] [lý do]",   "Đề xuất xử phạt (Ban/Kick) gửi lên nhóm Staff để Admin duyệt"],
    ],
    col_widths=[3.5, 4.5, 7.5]
)
add_note("💡  Tất cả lệnh Staff đều hỗ trợ cả Slash Command (/warn) và Prefix Command (!warn). "
         "Mỗi lệnh gửi DM cho người bị xử phạt và log vào kênh mod-log.", bg="E8F4FD", border="3498DB")

# ── 4.3 Admin ────────────────────────────────────────────────────────────────
add_heading("4.3 Lệnh Chỉ Dành Cho Admin", 2)
add_table(
    ["Lệnh", "Mô tả"],
    [
        ["/postrules",  "Xóa tin nhắn cũ và đăng lại toàn bộ nội quy vào kênh #rules chính thức"],
    ],
    col_widths=[4, 11.5]
)

doc.add_page_break()

# ── 4.4 /config ──────────────────────────────────────────────────────────────
add_heading("4.4 Nhóm Lệnh /config (Chỉ Admin)", 2)
add_para("Tất cả lệnh /config chỉ hiện với người gõ (ephemeral). "
         "Thay đổi có hiệu lực ngay, lưu vào data/config.json.", size=10, italic=True, color=C_GRAY)

add_heading("/config view", 3)
add_para("Xem toàn bộ cấu hình hiện tại kèm menu tương tác để điều hướng giữa các mục "
         "(Tổng quan / Ngưỡng Auto-Mod / Hệ thống phạt / Từ khóa / Domain).", size=10)

add_heading("/config set", 3)
add_table(
    ["Key", "Kiểu", "Mặc định", "Mô tả"],
    [
        ["bot_prefix",          "text",   "!",        "Ký tự prefix cho Prefix Commands"],
        ["guild_id",            "số",     "0",         "ID server (để sync slash command nhanh hơn)"],
        ["spam_threshold",      "số",     "5",         "Số tin nhắn giống nhau = spam"],
        ["spam_window",         "số",     "5",         "Cửa sổ thời gian spam (giây)"],
        ["max_emoji",           "số",     "15",        "Số emoji tối đa trong 1 tin nhắn"],
        ["max_mentions",        "số",     "5",         "Số @mention tối đa trong 1 tin nhắn"],
        ["mod_log_channel_id",  "số",     "0",         "ID kênh mod-log (thay bằng /config channel)"],
        ["rules_channel_id",    "số",     "0",         "ID kênh rules"],
        ["appeal_channel_id",   "số",     "0",         "ID kênh appeal"],
        ["staff_channel_id",    "số",     "0",         "ID kênh Staff để nhận đơn đề xuất phạt"],
        ["mod_role_id",         "số",     "0",         "ID role Moderator"],
        ["admin_role_id",       "số",     "0",         "ID role Admin"],
        ["muted_role_id",       "số",     "0",         "ID role Muted"],
    ],
    col_widths=[4, 1.5, 2, 8]
)

add_heading("/config channel", 3)
add_table(
    ["Chức năng", "Lệnh ví dụ", "Mô tả"],
    [
        ["mod_log",  "/config channel mod_log #mod-log",   "Kênh nhận log hành động mod"],
        ["rules",    "/config channel rules #rules",        "Kênh đăng nội quy"],
        ["appeal",   "/config channel appeal #appeal",      "Kênh nhận kháng cáo"],
        ["staff",    "/config channel staff #phòng-staff",  "Kênh nhận đề xuất xử phạt"],
    ],
    col_widths=[2.5, 6, 7]
)

add_heading("/config role", 3)
add_table(
    ["Chức năng", "Lệnh ví dụ", "Mô tả"],
    [
        ["moderator",  "/config role moderator @Mod",     "Role Moderator – dùng được lệnh Staff"],
        ["admin",      "/config role admin @Admin",        "Role Admin – dùng được /config"],
        ["muted",      "/config role muted @Muted",        "Role Muted – gán khi timeout"],
    ],
    col_widths=[2.5, 6, 7]
)

add_heading("/config warn", 3)
add_para("Thay đổi hình phạt cho warn thứ N:", size=10)
add_code("/config warn [số] [action] [phút]")
add_table(
    ["Action", "Ý nghĩa", "Ví dụ"],
    [
        ["warn_only",  "Chỉ cảnh cáo, không phạt thêm",    "/config warn 1 warn_only 0"],
        ["mute",       "Mute (timeout) trong X phút",       "/config warn 2 mute 60"],
        ["ban",        "Ban vĩnh viễn",                     "/config warn 5 ban 0"],
    ],
    col_widths=[2.5, 6.5, 6.5]
)

add_heading("/config addword & removeword", 3)
add_table(
    ["Danh sách", "Lệnh thêm", "Lệnh xóa"],
    [
        ["hate_keywords",     "/config addword hate_keywords <từ>",     "/config removeword hate_keywords <từ>"],
        ["politics_keywords", "/config addword politics_keywords <từ>",  "/config removeword politics_keywords <từ>"],
        ["suspicious_domains","/config addword suspicious_domains <domain>", "/config removeword suspicious_domains <domain>"],
        ["ad_patterns",       "/config addword ad_patterns <pattern>",   "/config removeword ad_patterns <pattern>"],
    ],
    col_widths=[3.5, 6, 6]
)

add_heading("/config reset", 3)
add_para("Reset một key về giá trị mặc định. Có thể reset: spam_threshold, spam_window, max_emoji, "
         "max_mentions, warn_punishments, hate_keywords, politics_keywords, suspicious_domains.", size=10)
add_code("/config reset spam_threshold")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  5. AUTO-MODERATION
# ════════════════════════════════════════════════════════════════════════════

add_heading("5. Hệ Thống Auto-Moderation", 1)
add_para("Bot tự động phát hiện và xử lý các vi phạm mà không cần staff can thiệp. "
         "Tất cả ngưỡng đều có thể chỉnh qua /config.", size=10)

add_table(
    ["Loại vi phạm", "Điều kiện phát hiện", "Hành động", "Warn?"],
    [
        ["Spam",              f"≥ spam_threshold tin nhắn giống nhau trong spam_window giây",
                              "Xóa tin nhắn",         "✅"],
        ["Link độc hại",      "Domain nằm trong suspicious_domains (IP grabber, Nitro giả...)",
                              "Xóa tin nhắn",         "✅"],
        ["Quảng cáo",         "Link khớp ad_patterns (discord.gg/, t.me/...)",
                              "Xóa tin nhắn",         "✅"],
        ["Ngôn từ thù ghét",  "Từ nằm trong hate_keywords",
                              "Xóa tin nhắn",         "✅"],
        ["Nội dung chính trị","Từ nằm trong politics_keywords",
                              "Xóa tin nhắn",         "✅"],
        ["Thông tin cá nhân", "Số điện thoại VN (0[3-9]xxxxxxxx hoặc +84...)",
                              "Xóa tin nhắn",         "✅"],
        ["Mention spam",      f"> max_mentions @mention trong 1 tin nhắn",
                              "Xóa tin nhắn",         "✅"],
        ["Emoji spam",        f"> max_emoji emoji trong 1 tin nhắn",
                              "Xóa tin nhắn",         "❌"],
    ],
    col_widths=[3.5, 5.5, 3, 1.5]
)

add_heading("Luồng xử lý khi phát hiện vi phạm", 2)
add_code("1. Xóa tin nhắn vi phạm")
add_code("2. Thêm warn vào warnings.json")
add_code("3. Kiểm tra tổng số warn → áp dụng hình phạt leo thang")
add_code("4. DM thông báo cho người vi phạm (kèm lý do + tổng warn)")
add_code("5. Gửi log chi tiết vào kênh mod-log")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  6. HỆ THỐNG XỬ PHẠT
# ════════════════════════════════════════════════════════════════════════════

add_heading("6. Hệ Thống Cảnh Cáo & Xử Phạt", 1)

add_heading("Mặc định leo thang theo số warn", 2)
add_table(
    ["Số Warn", "Hành động", "Thời gian", "Ghi chú"],
    [
        ["1", "⚠️ Cảnh cáo",    "—",       "DM + log"],
        ["2", "⚠️ Cảnh cáo",    "—",       "DM + log"],
        ["3", "🔇 Mute",        "1 giờ",   "Timeout Discord"],
        ["4", "🔇 Mute",        "2 giờ",   "Timeout Discord"],
        ["5", "🔇 Mute",        "7 giờ",   "Timeout Discord"],
        ["6", "🔇 Mute",        "14 giờ",  "Timeout Discord"],
        ["7", "🔇 Mute",        "28 giờ",  "Timeout Discord"],
        ["8", "🔇 Mute",        "56 giờ",  "Timeout Discord"],
        ["9", "🔇 Mute",        "72 giờ",  "Timeout Discord"],
        ["10+", "🔇 Mute + BAN","72 giờ",  "Mute 72h + Đề xuất BAN tự động lên kênh Staff để Admin duyệt"],
    ],
    col_widths=[2, 3.5, 3, 7]
)
add_note("💡  Toàn bộ mức phạt này có thể thay đổi qua /config warn. "
         "Ví dụ: /config warn 3 ban 0 → warn lần 3 sẽ ban luôn.", bg="E8F4FD", border="3498DB")

add_heading("Cách tính warn", 2)
add_bullet("Warn cộng dồn, không giới hạn")
add_bullet("Staff có thể xóa warn bằng /clearwarns @user")
add_bullet("Warn từ auto-mod và warn thủ công của Staff đều tính chung")

add_heading("Đề xuất xử phạt (Proposal System)", 2)
add_para("Để tránh lạm quyền, Moderator không được tự ý Ban/Kick thành viên mà phải gửi đề xuất lên nhóm Staff để Admin duyệt:", size=10)
add_bullet("Sử dụng lệnh: /guardbot propose <member> <ban/kick> <lý do> để tạo đơn đề xuất.")
add_bullet("Đơn đề xuất sẽ được gửi vào Staff Channel kèm theo nút ✅ Duyệt và ❌ Từ chối.")
add_bullet("Chỉ Admin hoặc chủ server mới có quyền nhấn duyệt hoặc từ chối đơn.")
add_bullet("Khi cảnh cáo thứ 10+ được kích hoạt (cả thủ công và tự động), bot sẽ tự động tạo đề xuất BAN và gửi vào nhóm Staff.")

add_heading("Kháng cáo", 2)
add_para("Khi bị xử phạt, thành viên có thể gõ /appeal để mở form kháng cáo:", size=10)
add_bullet("Điền: vi phạm bị xử phạt / giải trình / case ID")
add_bullet("Bot gửi kháng cáo vào kênh appeal kèm 2 nút: ✅ Chấp nhận / ❌ Từ chối")
add_bullet("Staff xử lý và bot tự DM kết quả cho người kháng cáo")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  7. EVENT LOGGING
# ════════════════════════════════════════════════════════════════════════════

add_heading("7. Hệ Thống Ghi Log Sự Kiện", 1)
add_para("Bot tự động ghi log các sự kiện quan trọng vào kênh mod-log:", size=10)

add_table(
    ["Sự kiện", "Thông tin ghi lại"],
    [
        ["📥 Thành viên tham gia",    "Tên, ID, ngày tạo tài khoản, tổng số thành viên"],
        ["📤 Thành viên rời server",  "Tên, ID, danh sách role đang có"],
        ["🗑️ Tin nhắn bị xóa",       "Tác giả, kênh, nội dung tin nhắn (tối đa 500 ký tự)"],
        ["✏️ Tin nhắn bị sửa",        "Tác giả, kênh, nội dung trước và sau, link đến tin nhắn"],
        ["🔇 Thành viên bị timeout",  "Tên, thời gian hết timeout"],
        ["🔊 Thành viên bỏ timeout",  "Tên thành viên"],
        ["⚖️ Lệnh mod thủ công",      "Hành động, người bị phạt, moderator, lý do, case ID"],
        ["🤖 Auto-Mod phát hiện",     "Loại vi phạm, người dùng, nội dung vi phạm"],
    ],
    col_widths=[4.5, 11]
)

# ════════════════════════════════════════════════════════════════════════════
#  8. DỮ LIỆU & LƯU TRỮ
# ════════════════════════════════════════════════════════════════════════════

add_heading("8. Dữ Liệu & Lưu Trữ", 1)

add_heading("warnings.json – Lịch sử cảnh cáo", 2)
add_code('{')
add_code('  "USER_ID": [')
add_code('    {')
add_code('      "id": 1,              // Số thứ tự warn')
add_code('      "reason": "Spam",     // Lý do')
add_code('      "moderator_id": 123,  // ID mod/bot xử lý')
add_code('      "guild_id": 456,      // ID server')
add_code('      "timestamp": "2024-..." // Thời điểm')
add_code('    }')
add_code('  ]')
add_code('}')

add_heading("cases.json – Lịch sử xử lý", 2)
add_code('{')
add_code('  "1": {')
add_code('    "id": 1,              // Case ID')
add_code('    "action": "WARN",     // WARN / MUTE / BAN / KICK / AUTO-WARN...')
add_code('    "target_id": ...,     // ID người bị xử lý')
add_code('    "target_tag": "...",  // Tên#0000')
add_code('    "mod_id": ...,        // ID moderator')
add_code('    "reason": "...",      // Lý do')
add_code('    "duration_min": 60,   // Thời gian (nếu mute)')
add_code('    "timestamp": "..."    // Thời điểm')
add_code('  }')
add_code('}')

add_heading("config.json – Cấu hình runtime", 2)
add_para("Tự động tạo khi chạy bot lần đầu. Có thể xóa file này để reset về mặc định.", size=10, italic=True, color=C_GRAY)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  9. QUICK REFERENCE
# ════════════════════════════════════════════════════════════════════════════

add_heading("9. Quick Reference – Tra Cứu Nhanh", 1)

add_table(
    ["Tình huống", "Lệnh cần dùng"],
    [
        ["Setup lần đầu (kênh)",           "/config channel mod_log / rules / appeal"],
        ["Setup lần đầu (role)",            "/config role moderator / admin / muted"],
        ["Cảnh cáo ai đó",                  "/warn @user [lý do]"],
        ["Mute ai đó 30 phút",              "/mute @user 30 [lý do]"],
        ["Bỏ mute",                         "/unmute @user"],
        ["Kick khỏi server",                "/kick @user [lý do]"],
        ["Ban ai đó",                       "/ban @user [lý do]"],
        ["Bỏ ban theo ID",                  "/unban 123456789 [lý do]"],
        ["Xem warn của ai",                 "/warns @user"],
        ["Xóa toàn bộ warn",               "/clearwarns @user"],
        ["Xóa 50 tin nhắn",                "/purge 50"],
        ["Xem case #5",                     "/case 5"],
        ["Thêm từ cấm",                     "/config addword hate_keywords <từ>"],
        ["Xóa từ cấm",                      "/config removeword hate_keywords <từ>"],
        ["Thêm domain xấu",                "/config addword suspicious_domains <domain>"],
        ["Đổi ngưỡng spam",                "/config set spam_threshold 3"],
        ["Đổi hình phạt warn 3",           "/config warn 3 ban 0"],
        ["Đăng nội quy vào kênh #rules",   "/postrules"],
        ["Xem & chỉnh toàn bộ config",     "/config view"],
    ],
    col_widths=[6, 9.5]
)

# ─── Footer ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("Discord Moderation Bot  •  Tài liệu nội bộ  •  Không phân phối")
fr.font.size      = Pt(8)
fr.font.color.rgb = C_GRAY

# ─── Lưu file ────────────────────────────────────────────────────────────────
output = r"c:\py\discord-bot\HuongDan_Bot_Discord.docx"
doc.save(output)
print(f"Docx generated successfully: {output}")
